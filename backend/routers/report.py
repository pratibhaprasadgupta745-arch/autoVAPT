from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from database import SessionLocal
from models import Scan, Vulnerability

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document
from openpyxl import Workbook

router = APIRouter(prefix="/report", tags=["report"])


# ---------------- DB ----------------
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ---------------- COMMON DATA ----------------
def get_scan_data(scan_id, db):
    scan = db.query(Scan).filter(Scan.id == scan_id).first()
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")

    vulns = db.query(Vulnerability).filter(
        Vulnerability.scan_id == scan_id
    ).all()

    return scan, vulns


# ---------------- PDF ----------------
@router.get("/pdf/{scan_id}")
def generate_pdf(scan_id: int, db: Session = Depends(get_db)):

    scan, vulns = get_scan_data(scan_id, db)

    file_path = f"report_{scan_id}.pdf"

    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()
    elements = []

    # LOGO STYLE HEADER
    elements.append(Paragraph("🔐 AutoVAPT Security Report", styles["Title"]))
    elements.append(Spacer(1, 15))

    elements.append(Paragraph(f"Target: {scan.target}", styles["Normal"]))
    elements.append(Paragraph(f"Status: {scan.status}", styles["Normal"]))
    elements.append(Paragraph(f"Risk Score: {scan.risk_score}", styles["Normal"]))
    elements.append(Spacer(1, 20))

    elements.append(Paragraph("Vulnerabilities", styles["Heading2"]))
    elements.append(Spacer(1, 10))

    for v in vulns:
        elements.append(Paragraph(f"{v.name} | {v.severity}", styles["Normal"]))
        elements.append(Paragraph(f"Fix: {v.fix}", styles["Normal"]))
        elements.append(Spacer(1, 10))

    doc.build(elements)

    return {"file": file_path}


# ---------------- DOCX ----------------
@router.get("/docx/{scan_id}")
def generate_docx(scan_id: int, db: Session = Depends(get_db)):

    scan, vulns = get_scan_data(scan_id, db)

    file_path = f"report_{scan_id}.docx"

    doc = Document()
    doc.add_heading("AutoVAPT Security Report", 0)

    doc.add_paragraph(f"Target: {scan.target}")
    doc.add_paragraph(f"Risk Score: {scan.risk_score}")

    for v in vulns:
        doc.add_heading(v.name, level=2)
        doc.add_paragraph(f"Severity: {v.severity}")
        doc.add_paragraph(f"Fix: {v.fix}")

    doc.save(file_path)

    return {"file": file_path}


# ---------------- EXCEL ----------------
@router.get("/excel/{scan_id}")
def generate_excel(scan_id: int, db: Session = Depends(get_db)):

    scan, vulns = get_scan_data(scan_id, db)

    file_path = f"report_{scan_id}.xlsx"

    wb = Workbook()
    ws = wb.active

    ws.append(["Name", "Severity", "Path", "Fix"])

    for v in vulns:
        ws.append([v.name, v.severity, v.path, v.fix])

    wb.save(file_path)

    return {"file": file_path}